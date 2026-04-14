"""
DOCX Generator Template for Review Papers (Python)

USAGE:
  1. pip install python-docx
  2. Replace the placeholder content below with your actual paper content
  3. python generate_docx_template.py

This template follows the review_paper_format_guide.md formatting rules:
  - Times New Roman throughout
  - 24pt Bold centered title
  - 3-column borderless author table
  - 7-column comparison table
  - US Letter size (8.5 x 11 inches)
"""

import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── CONFIGURATION ────────────────────────────────────────────────────────────
CONFIG = {
    'title': '[YOUR PAPER TITLE HERE]',
    'outputFile': 'REVIEW_PAPER.docx',
    'authors': [
        {'name': 'Student A', 'dept': 'Department of Computer Science', 'uni': 'University Name, City, State', 'email': 'studentA@university.edu'},
        {'name': 'Student B', 'dept': 'Department of Computer Science', 'uni': 'University Name, City, State', 'email': 'studentB@university.edu'},
        {'name': 'Professor Name', 'designation': 'Assistant Professor', 'dept': 'Department of Computer Science', 'uni': 'University Name, City, State', 'email': 'professor@university.edu'},
    ]
}

# ─── CONTENT (REPLACE THESE WITH YOUR ACTUAL CONTENT) ─────────────────────────
abstractText = '[YOUR ABSTRACT TEXT HERE - 200-300 words covering all 6 required points]'
keywordsText = 'Keyword One, Keyword Two, Keyword Three, Keyword Four, Keyword Five'

introParas = [
    '[Introduction paragraph 1]',
    '[Introduction paragraph 2]',
    '[Introduction paragraph 3]',
]

objectiveText = '[Your objective paragraph here - 5-8 sentences, third person]'

litReview = [
    # {'text': 'Paragraph text here...', 'citation': '[1]'},
    # ... repeat for all 30 papers
]

compRows = [
    # ['1', 'Paper Title', 'Author et al. (Year)', 'Year', 'Objective', 'Methodology', 'Result'],
    # ... repeat for 5 papers
]

conclusionParas = [
    '[Conclusion paragraph 1 - summary]',
    '[Conclusion paragraph 2 - challenges]',
    '[Conclusion paragraph 3 - forward-looking]',
]

futureText = '[Future scope paragraph - ~150 words]'

references = [
    # '[1] Author, A. (Year). Title. Journal, Vol(Issue), pages.',
    # ... repeat for 30
]

# ─── HELPERS ──────────────────────────────────────────────────────────────────
doc = Document()

# Set US Letter size and 1-inch margins
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

def set_font(run, size=12, bold=False, italic=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def add_heading(text, level=1, size=16):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, size=size, bold=True)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)

def add_paragraph(text, size=12, bold=False, justify=True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run, size, bold)
    p.paragraph_format.space_after = Pt(6)
    return p

# ─── BUILD DOCUMENT ───────────────────────────────────────────────────────────

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(CONFIG['title'])
set_font(run, size=24, bold=True)
p.paragraph_format.space_after = Pt(24)

# Author Block
table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
cells = table.rows[0].cells
for i, author in enumerate(CONFIG['authors']):
    cell = cells[i]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(author['name'] + "\n")
    set_font(run, size=12, bold=True)

    if 'designation' in author:
        run = p.add_run(author['designation'] + "\n")
        set_font(run, size=12)
        
    run = p.add_run(author['dept'] + "\n" + author['uni'] + "\n")
    set_font(run, size=12)
    
    run = p.add_run(author['email'])
    set_font(run, size=12)
    run.font.underline = True
    run.font.color.rgb = RGBColor(0, 0, 255) # blue link

doc.add_paragraph() # spacing

# Abstract
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run_label = p.add_run("Abstract\u2014")
set_font(run_label, size=12, bold=True)
run_text = p.add_run(abstractText)
set_font(run_text, size=12)

# Keywords
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(18)
run_label = p.add_run("Keywords\u2014")
set_font(run_label, size=12, bold=True)
run_text = p.add_run(keywordsText)
set_font(run_text, size=12)

# I. Introduction
add_heading("I. Introduction")
for para in introParas:
    add_paragraph(para)

# II. Objective
add_heading("II. Objective")
add_paragraph(objectiveText)

# III. Literature Review
add_heading("III. Literature Review")
for item in litReview:
    p = add_paragraph(item['text'] + " ")
    run_cite = p.add_run(item['citation'])
    set_font(run_cite, size=12, bold=True)

# IV. Comparison
add_heading("IV. Comparison of 30 Published Research Papers")
add_paragraph("[Introductory paragraph explaining why comparison is needed]")

p_caption = doc.add_paragraph()
run_caption = p_caption.add_run("Table 1. Comparative Summary of Selected Published Research Papers")
set_font(run_caption, size=9, bold=True)

if compRows:
    comp_table = doc.add_table(rows=1, cols=7)
    comp_table.style = 'Grid Table 1 Light' # gives basic borders
    hdr_cells = comp_table.rows[0].cells
    headers = ['S.No', 'Title of the RP', 'Author Name', 'Year', 'Objective', 'Methodology', 'Conclusion / Result']
    for i, hc in enumerate(headers):
        hdr_cells[i].text = hc
        set_font(hdr_cells[i].paragraphs[0].runs[0], size=10, bold=True)
        
    for row_data in compRows:
        row_cells = comp_table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
            set_font(row_cells[i].paragraphs[0].runs[0], size=10)

# V. Conclusion
add_heading("V. Conclusion")
for para in conclusionParas:
    add_paragraph(para)

# VI. Future Scope
add_heading("VI. Future Scope")
add_paragraph(futureText)

# References
add_heading("References")
for ref in references:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    run = p.add_run(ref)
    set_font(run, size=10)

doc.save(CONFIG['outputFile'])
print(f"Done! Generated {CONFIG['outputFile']}")
